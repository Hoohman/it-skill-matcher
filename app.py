from flask import Flask, render_template, request, jsonify, send_file
from werkzeug.exceptions import RequestEntityTooLarge
from docx import Document
import fitz
import re
import json
import io

app = Flask(__name__)

app.config['MAX_CONTENT_LENGTH'] = 5 * 1024 * 1024 # 5MB limit

ALLOWED_EXTENSIONS = {'pdf', 'docx'}

# --- Функции обработки файлов ---
def load_data_from_json(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            return data['IT_SKILLS'], data['ROLE_SKILLS']
    except Exception as e:
        print(f"Ошибка загрузки {file_path}: {e}")
        # Возвращаем значения по умолчанию
        default_it = ["ошибка"]
        default_roles = {"Default": ["ошибка"]}
        return default_it, default_roles
    
IT_SKILLS, ROLE_SKILLS = load_data_from_json('it_skills.json')

def extract_text_from_docx(uploaded_file):
    doc = Document(uploaded_file)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_pdf(uploaded_file):
    doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def extract_skills(text):
    text = text.lower()
    text = re.sub(r'[^a-zа-я0-9\s\/\+\-]', ' ', text)
    found_skills = [skill for skill in IT_SKILLS if skill.lower() in text]
    return list(set(found_skills))

def compare_skills(resume_skills, vacancy_skills):
    # vacancy_skills теперь будет словарем с приоритетами
    critical = set(vacancy_skills.get('critical', []))
    important = set(vacancy_skills.get('important', []))
    optional = set(vacancy_skills.get('optional', []))
    
    all_required = critical.union(important).union(optional)
    
    # Базовые совпадения
    matched_all = list(set(resume_skills) & all_required)
    missing_all = list(all_required - set(resume_skills))
    
    # Совпадения по приоритетам
    matched_critical = list(set(resume_skills) & critical)
    matched_important = list(set(resume_skills) & important)
    matched_optional = list(set(resume_skills) & optional)
    
    # Расчет взвешенного процента (критические x3, важные x2, дополнительные x1)
    total_weight = len(critical)*3 + len(important)*2 + len(optional)*1
    if total_weight == 0:
        return {}, [], [], 0
    
    match_weight = (len(matched_critical)*3 + len(matched_important)*2 + len(matched_optional)*1)
    match_percent = round(100 * match_weight / total_weight, 1)
    
    return {
        'critical': {
            'matched': matched_critical,
            'missing': list(critical - set(resume_skills))
        },
        'important': {
            'matched': matched_important,
            'missing': list(important - set(resume_skills))
        },
        'optional': {
            'matched': matched_optional,
            'missing': list(optional - set(resume_skills))
        }
    }, matched_all, missing_all, match_percent

def match_roles(user_skills, role_skills_dict, vacancy_skills=None, min_match=0):
    role_matches = {}
    for role, skills_required in role_skills_dict.items():
        matched = set(user_skills) & set(skills_required)
        
        if vacancy_skills:
            relevant_matched = matched & set(vacancy_skills)
        else:
            relevant_matched = matched

        match_percent = round(100 * len(relevant_matched) / max(len(skills_required), 1), 1)
        if match_percent >= min_match:
            role_matches[role] = {
                "matched_skills": list(matched),
                "relevant_matched_skills": list(relevant_matched),
                "match_percent": match_percent
            }
    return role_matches

def extract_name(text):
    lines = [line.strip() for line in text.strip().split("\n") if line.strip()]
    ignore_keywords = {"резюме", "портрет", "фото"}
    patronymics_keywords = {"вич", "вна"}

    def has_patronymic(word):
        return any(word.lower().endswith(p) for p in patronymics_keywords)

    for line in lines:
        # Обработка строк с "ФИО:", "Имя:", "Name:"
        if re.search(r"(ФИО|Имя|Name)\s*[:\-]", line, re.IGNORECASE):
            possible_name = re.split(r"[:\-]", line, maxsplit=1)[1].strip()
            # Удаляем возможные остатки меток (если разделитель был пробелом)
            possible_name = re.sub(r'^(ФИО|Имя|Name)\s*', '', possible_name, flags=re.IGNORECASE).strip()
            if 2 <= len(possible_name.split()) <= 4:
                return possible_name

        words = line.split()
        # Проверка на отчество и подходящую длину
        if any(has_patronymic(word) for word in words) and 2 <= len(words) <= 4:
            return " ".join(words)

        # Проверка на все слова с заглавной буквы
        if all(w.istitle() or w.isupper() for w in words) and 2 <= len(words) <= 4 and not any(word.lower() in ignore_keywords for word in words):
            return " ".join(words)

    return "Имя не найдено"

def extract_email(text):
    # Ищем стандартные форматы email
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, text)
    return emails[0] if emails else "не указан в анкете"

def extract_phone(text):
    # Ищем российские номера в различных форматах
    phone_pattern = r'(?:\+7|8)[\s\-]?\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{2}[\s\-]?\d{2}'
    phones = re.findall(phone_pattern, text)

    # Нормализуем номер
    if phones:
        phone = phones[0]
        # Удаляем все нецифровые символы
        phone = re.sub(r'[^\d]', '', phone)
        # Приводим к формату +7XXXXXXXXXX
        if phone.startswith('8'):
            phone = '+7' + phone[1:]
        elif phone.startswith('7'):
            phone = '+' + phone
        return phone
    return "не указан в анкете"

def create_report_docx(name, contacts, match_percent, matched, missing, role_results):
    doc = Document()
    doc.add_heading(f'Анализ резюме: {name}', 0)

    doc.add_heading('Контактная информация', level=1)
    doc.add_paragraph(f'Email: {contacts.get("email", "не указан")}')
    doc.add_paragraph(f'Телефон: {contacts.get("phone", "не указан")}')

    doc.add_heading('Общая оценка совпадения с вакансией', level=1)
    doc.add_paragraph(f'Совпадение по навыкам: {match_percent}%')

    doc.add_heading('Совпавшие навыки с вакансией', level=1)
    doc.add_paragraph(", ".join(matched) if matched else "Совпадений нет")

    doc.add_heading('Недостающие навыки', level=1)
    doc.add_paragraph(", ".join(missing) if missing else "Нет пробелов")

    doc.add_heading('Соответствие профессиональным ролям', level=1)
    for role, info in role_results.items():
        doc.add_heading(f'Роль: {role} ({info["match_percent"]}%)', level=2)
        doc.add_paragraph(f'Совпавшие навыки: {", ".join(info["matched_skills"])}')
        if "relevant_matched_skills" in info:
            doc.add_paragraph(f'Навыки, совпавшие с вакансией: {", ".join(info["relevant_matched_skills"])}')

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --- Маршруты Flask ---
@app.errorhandler(RequestEntityTooLarge)
def handle_file_too_large(e):
    return jsonify({"error": "Размер файла превышает 5 МБ"}), 413

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_files():
    resume_file = request.files.get('resume')
    vacancy_source = request.form.get('vacancy_source', 'file')
    
    if not resume_file:
        return jsonify({"error": "Пожалуйста, загрузите резюме"}), 400
        
    # Обработка резюме
    resume_text = extract_text_from_docx(resume_file) if resume_file.filename.endswith(".docx") else extract_text_from_pdf(resume_file)
    resume_skills = extract_skills(resume_text)
    name = extract_name(resume_text)
    
    # Определяем навыки вакансии
    if vacancy_source == 'file':
        vacancy_file = request.files.get('vacancy')
        if not vacancy_file:
            return jsonify({"error": "Пожалуйста, загрузите файл вакансии"}), 400
        vacancy_text = extract_text_from_docx(vacancy_file) if vacancy_file.filename.endswith(".docx") else extract_text_from_pdf(vacancy_file)
        vacancy_skills = {
            'critical': extract_skills(vacancy_text),  # Для файла все навыки считаем критическими
            'important': [],
            'optional': []
        }
    else:
        # Получаем навыки из формы
        try:
            vacancy_skills = {
                'critical': json.loads(request.form.get('critical_skills', '[]')),
                'important': json.loads(request.form.get('important_skills', '[]')),
                'optional': json.loads(request.form.get('optional_skills', '[]'))
            }
        except json.JSONDecodeError:
            return jsonify({"error": "Ошибка в формате навыков"}), 400
    
    # Контактная информация
    email = extract_email(resume_text)
    phone = extract_phone(resume_text)

    # Сравниваем навыки
    skills_by_priority, matched, missing, match_percent = compare_skills(resume_skills, vacancy_skills)
    role_results = match_roles(resume_skills, ROLE_SKILLS, None, 1)

    return jsonify({
        "name": name,
        "contacts": {"email": email, "phone": phone},
        "skills_by_priority": skills_by_priority,
        "matched": matched,
        "missing": missing,
        "match_percent": match_percent,
        "role_results": role_results,
        "resume_text": resume_text
    })
    
@app.route('/api/skills', methods=['GET'])
def get_skills_api():
    return jsonify({
        'IT_SKILLS': IT_SKILLS,
        'ROLE_SKILLS': ROLE_SKILLS
    })
    
@app.route('/download-report', methods=['POST'])
def download_report():
    data = request.get_json()
    name = data.get("name")
    contacts = data.get("contacts")
    match_percent = data.get("match_percent")
    matched = data.get("matched")
    missing = data.get("missing")
    role_results = data.get("role_results")

    doc_buffer = create_report_docx(name, contacts, match_percent, matched, missing, role_results)

    return send_file(
        doc_buffer,
        as_attachment=True,
        download_name=f'Результаты_{name}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    app.run(debug=False)